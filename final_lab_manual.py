#!/usr/bin/env python3
"""
Final Lab Manual Generator for Python Notebooks
Creates a professional Word document with code and explanations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def set_font(run, font_name="Consolas", font_size=10, bold=False, color=None):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)
    
    # Section heading style
    section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Arial'
    section_style.font.size = Pt(16)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(204, 51, 0)
    section_style.paragraph_format.space_before = Pt(20)
    section_style.paragraph_format.space_after = Pt(10)
    
    # Code style
    code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Explanation style
    explanation_style = styles.add_style('Explanation', WD_STYLE_TYPE.PARAGRAPH)
    explanation_style.font.name = 'Arial'
    explanation_style.font.size = Pt(11)
    explanation_style.paragraph_format.space_after = Pt(10)
    explanation_style.paragraph_format.line_spacing = 1.15

def add_code_block(doc, code, title="Code:"):
    """Add a formatted code block"""
    # Add code title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    set_font(title_run, "Arial", 11, bold=True, color=(0, 51, 102))
    
    # Add code content
    code_para = doc.add_paragraph(style='CodeBlock')
    code_run = code_para.add_run(code)
    set_font(code_run, "Consolas", 10, color=(51, 51, 51))

def create_lab_manual():
    """Create the complete lab manual document"""
    print("Creating Python Lab Manual...")
    
    # Create document
    doc = Document()
    create_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_paragraph(style='CustomTitle')
    title.add_run("PYTHON PROGRAMMING\nLAB MANUAL")
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Fundamentals of Python Data Types and Operations")
    set_font(subtitle_run, "Arial", 14, color=(102, 102, 102))
    
    # Course info
    doc.add_paragraph()
    course_info = doc.add_paragraph()
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_run = course_info.add_run("Course: Python Programming\nDate: May 30, 2025")
    set_font(course_run, "Arial", 12, color=(102, 102, 102))
    
    add_page_break(doc)
    
    # Table of Contents
    toc_title = doc.add_paragraph(style='CustomTitle')
    toc_title.add_run("TABLE OF CONTENTS")
    
    toc_items = [
        ("1. Basic Data Types and Type Conversion", "3"),
        ("2. String Operations", "5"),
        ("3. Lists Operations", "7"),
        ("4. Tuples Operations", "9"),
        ("5. Sets Operations", "11"),
        ("6. Summary and Conclusion", "13")
    ]
    
    for item, page in toc_items:
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run(f"{item}{'.' * (50 - len(item))}{page}")
        set_font(toc_run, "Arial", 11)
    
    add_page_break(doc)
    
    # Lab Content
    labs = [
        {
            "title": "LAB 1: BASIC DATA TYPES AND TYPE CONVERSION",
            "filename": "start.ipynb",
            "explanation": """
This lab introduces Python's fundamental data types including integers, floats, complex numbers, 
strings, booleans, None type, and lists. Students will learn to check variable types using the 
type() function and perform type conversions between different data types. The lab also covers 
basic user input handling and arithmetic operations.

Learning Objectives:
• Understand Python's basic data types
• Learn type checking with type() function
• Practice type conversion methods
• Handle user input and perform basic operations
            """,
            "code": """# Basic Data Types
a = 5 
print(a)
print(type(a))

b = 3j
print(b)
print(type(b))

c = 222.22
print(c)
print(type(c))

d = "hello world"
print(d)
print(type(d))

e = True
print(e)
print(type(e))

f = False
print(f)
print(type(f))

g = None
print(g)
print(type(g))

h = [1, 2, 3, 4, 5]
print(h)
print(type(h))

print('\\n')

# Type Conversion
d = int(c)
print(d)
print(type(d))

e = float(a)
print(e)    
print(type(e))

print('\\n')

cc = complex(a)
print(cc)
print(type(cc))

# User Input and Basic Operations
aa = input("Enter a number: ")
ss = int(aa)
print(ss+5)

print('\\n')

bb = input("Enter a number: ")
kk = int(bb)
print(kk%5)

print('\\n')

dd = input("Enter a number: ")
mm = int(dd)
print(mm/5)

ff = input("Enter a number: ")
nn = int(ff)
print(nn*5)""",
            "summary": """
Key Takeaways:
• Python supports multiple data types: int, float, complex, str, bool, NoneType, list
• Use type() function to check variable types
• Type conversion functions: int(), float(), complex(), str()
• input() function returns string by default, requiring conversion for numerical operations
• Python is dynamically typed - variables can change types during execution
            """
        },
        {
            "title": "LAB 2: STRING OPERATIONS",
            "filename": "second.ipynb",
            "explanation": """
This lab focuses on string manipulation in Python. Students will learn various string methods 
for case conversion, text cleaning, replacement, splitting, and counting occurrences. These 
operations are fundamental for text processing and data cleaning tasks.

Learning Objectives:
• Master string case conversion methods
• Learn string manipulation techniques
• Practice text processing operations
• Understand string method chaining
            """,
            "code": """# Getting User Input
a = input("Enter a String: ")
print(a)
b = input("Enter a String: ")
print(b)

# String Case Conversion Methods
print(a.upper())

print(a.lower())

print(a.capitalize())

print(a.strip())

# String Manipulation Methods
print(a)

print(a.replace("a", "b"))

print(a.split("a"))

print(a.split(" "))

print(a.count("a"))""",
            "summary": """
Key Takeaways:
• String methods for case conversion: upper(), lower(), capitalize(), strip()
• String manipulation methods: replace(), split(), count()
• These methods return new strings (strings are immutable)
• strip() removes whitespace from beginning and end
• split() creates lists from strings based on delimiters
            """
        },
        {
            "title": "LAB 3: LISTS OPERATIONS",
            "filename": "list.ipynb",
            "explanation": """
This lab covers Python lists, which are ordered, mutable collections that can hold different 
data types. Students will learn list creation, indexing, slicing, concatenation, and various 
list methods essential for data manipulation.

Learning Objectives:
• Understand list creation and properties
• Master list indexing and slicing
• Learn list modification methods
• Practice list concatenation and extension
            """,
            "code": """# Creating Lists
my_list = ['orange', 'mango', 'banana', 'apple']
print(my_list)
print(type(my_list))

# List Slicing
print(my_list[3:5])

# Creating and Concatenating Lists
my_list2 = ['two','three']

print(*(my_list + my_list2), sep ='\\n')

# Adding Elements to Lists
my_list2.append('five')

my_list2.insert(0, 'one')
my_list2.insert(1, 'two')

# List Methods - Count and Index
print (my_list2.count('five'))
print (my_list2.count('one'))
print (my_list2.count('two'))

print (my_list2.index('five'))

# Extending Lists
my_list.extend(my_list2)
print(my_list)""",
            "summary": """
Key Takeaways:
• Lists are ordered, mutable collections enclosed in square brackets
• List slicing syntax: list[start:end] (end is exclusive)
• Methods for adding elements: append(), insert(), extend()
• Methods for finding elements: count(), index()
• Lists can be concatenated with + operator
• Use * operator with print() for unpacking list elements
            """
        },
        {
            "title": "LAB 4: TUPLES OPERATIONS",
            "filename": "tuples.ipynb",
            "explanation": """
This lab covers Python tuples, which are ordered, immutable collections. Unlike lists, tuples 
cannot be modified after creation, making them useful for storing data that shouldn't change. 
Students will learn tuple operations and conversions.

Learning Objectives:
• Understand tuple immutability and use cases
• Learn tuple creation and accessing methods
• Practice tuple methods and operations
• Master conversion between tuples and lists
            """,
            "code": """# Tuple Basics
my_tuple = ('orange', 'mango', 'banana', 'apple')
print("\\n--- Tuple Basics ---")
print(my_tuple)
print(type(my_tuple))
print(len(my_tuple))

# Accessing Tuple Elements
print("\\n--- Accessing Elements ---")
print(my_tuple[0])
print(my_tuple[-1])
print(my_tuple[1:3])

# Modifying Tuples (Converting to List and Back)
duplicat_list = list(my_tuple)

duplicat_list.append('kiwi')
duplicat_list.insert(0, 'xyz')

updated_tuple = tuple(duplicat_list)
print(updated_tuple)

# Tuple Methods
print("\\n--- Tuple Methods ---")
repeated_tuple = (1, 2, 3, 1, 2, 1)
print(repeated_tuple.count(1))
print(repeated_tuple.index(2))

# Tuple Concatenation
print("\\n--- Tuple Concatenation ---")
tuple1 = (1, 2, 3)
tuple2 = ('a', 'b', 'c')
combined = tuple1 + tuple2
print(combined)

# Nested Tuples
print("\\n--- Nested Tuples ---")
nested = ((1, 2), ('a', 'b'), (True, False))
print(nested)
print(nested[1][0])

# Type Conversions
print("\\n--- Conversions ---")
list_to_convert = ['one', 'two', 'three']

converted_tuple = tuple(list_to_convert)
print(f"List {list_to_convert} converted to tuple: {converted_tuple}")
back_to_list = list(converted_tuple)
print(f"Tuple converted back to list: {back_to_list}")

# Printing Tuple Elements
print("\\n--- Printing Tuple Elements Line by Line ---")
print(*my_tuple, sep='\\n')""",
            "summary": """
Key Takeaways:
• Tuples are ordered, immutable collections enclosed in parentheses
• Cannot modify tuples directly - convert to list for modifications
• Limited methods: count() and index()
• Support concatenation with + operator
• Can contain nested structures
• Useful for fixed data that shouldn't change
• Convert between tuple() and list() functions
            """
        },
        {
            "title": "LAB 5: SETS OPERATIONS",
            "filename": "set.ipynb",
            "explanation": """
This lab covers Python sets, which are unordered collections of unique elements. Sets are 
useful for removing duplicates, membership testing, and mathematical set operations. Students 
will also learn about frozensets.

Learning Objectives:
• Understand set properties and unique element constraint
• Learn set creation and modification methods
• Master mathematical set operations
• Practice with frozensets for immutable collections
            """,
            "code": """# Creating Sets
print("\\n--- Creating Sets ---")
set1 = {1, 2, 3, 4}
set2 = set(["apple", "banana", "cherry"])
empty_set = set()
print(f"Set1: {set1}")
print(f"Set2: {set2}")
print(f"Empty Set: {empty_set}")

# Adding Elements to Sets
print("\\n--- Adding Elements ---")
set1.add(5)
set1.update([6, 7, 8])
print(f"Set1 after adding elements: {set1}")

# Removing Elements from Sets
print("\\n--- Removing Elements ---")
set1.discard(8)
set1.remove(7)
print(f"Set1 after removing elements: {set1}")

# Accessing Elements in Sets
print("\\n--- Accessing Elements ---")
for item in set2:
    print(item)

# Set Operations
print("\\n--- Set Operations ---")
set_a = {1, 2, 3, 4}
set_b = {3, 4, 5, 6}
print(f"Union: {set_a | set_b}")
print(f"Intersection: {set_a & set_b}")
print(f"Difference (A - B): {set_a - set_b}")

# Clearing Sets
print("\\n--- Clearing a Set ---")
set1.clear()
print(f"Set1 after clearing: {set1}")

# Frozensets (Immutable Sets)
print("\\n--- Frozenset (Immutable Set) ---")
frozen = frozenset([1, 2, 3])
print(f"Frozenset: {frozen}")""",
            "summary": """
Key Takeaways:
• Sets store unique elements only (automatically remove duplicates)
• Created with curly braces {} or set() function
• Methods for adding: add(), update()
• Methods for removing: discard(), remove(), clear()
• Set operations: union (|), intersection (&), difference (-)
• Sets are unordered - no indexing
• Frozensets are immutable versions of sets
• Useful for membership testing and mathematical operations
            """
        }
    ]
    
    # Add each lab section
    for i, lab in enumerate(labs, 1):
        # Section title
        section_title = doc.add_paragraph(style='SectionHeading')
        section_title.add_run(lab['title'])
        
        # Source file info
        source_para = doc.add_paragraph()
        source_run = source_para.add_run(f"Source File: {lab['filename']}")
        set_font(source_run, "Arial", 10, color=(102, 102, 102))
        
        # Explanation
        explanation_para = doc.add_paragraph(style='Explanation')
        explanation_para.add_run(lab['explanation'].strip())
        
        # Code section
        add_code_block(doc, lab['code'], "Python Code:")
        
        # Summary
        summary_title = doc.add_paragraph()
        summary_run = summary_title.add_run("Summary and Analysis:")
        set_font(summary_run, "Arial", 12, bold=True, color=(0, 102, 51))
        
        summary_para = doc.add_paragraph(style='Explanation')
        summary_para.add_run(lab['summary'].strip())
        
        # Add page break except for last lab
        if i < len(labs):
            add_page_break(doc)
    
    # Final Summary Page
    add_page_break(doc)
    final_title = doc.add_paragraph(style='CustomTitle')
    final_title.add_run("COURSE SUMMARY AND CONCLUSION")
    
    conclusion_text = """
This comprehensive lab manual covered the fundamental data structures and operations in Python programming. 
Through hands-on exercises, students have gained practical experience with:

CORE CONCEPTS MASTERED:
• Data Types: Understanding integers, floats, complex numbers, strings, booleans, and None type
• Type Conversion: Converting between different data types using built-in functions
• String Operations: Text manipulation, case conversion, and string methods
• Lists: Mutable sequences with indexing, slicing, and modification methods
• Tuples: Immutable sequences for fixed data storage
• Sets: Unique collections with mathematical operations

PROGRAMMING SKILLS DEVELOPED:
• Variable declaration and type checking
• User input handling and validation
• Data structure selection based on requirements
• Method chaining and functional programming concepts
• Problem-solving with appropriate data structures

PRACTICAL APPLICATIONS:
These fundamental concepts form the foundation for:
• Data analysis and manipulation
• Web development and API design
• Scientific computing and research
• Database operations and data modeling
• Algorithm implementation and optimization

NEXT STEPS:
Students should continue practicing these concepts and explore:
• Advanced data structures (dictionaries, nested structures)
• File handling and data persistence
• Object-oriented programming principles
• Error handling and debugging techniques
• Third-party libraries and frameworks

This lab manual serves as a reference guide for future Python programming endeavors.
    """
    
    conclusion_para = doc.add_paragraph(style='Explanation')
    conclusion_para.add_run(conclusion_text.strip())
    
    # Save document
    filename = "/media/notsaaadkhan/DATA/UniPython/Python_Lab_Manual.docx"
    doc.save(filename)
    print(f"✅ Lab manual saved as: {filename}")
    
    return filename

if __name__ == "__main__":
    try:
        filename = create_lab_manual()
        print(f"\n🎉 SUCCESS! Your lab manual has been created: {filename}")
        print("\nFeatures included:")
        print("✓ Professional title page")
        print("✓ Table of contents")
        print("✓ Detailed explanations for each lab")
        print("✓ Formatted code blocks")
        print("✓ Learning objectives and summaries")
        print("✓ Course conclusion and next steps")
        print("✓ Professional styling and formatting")
        
    except Exception as e:
        print(f"❌ Error creating lab manual: {e}")
        import traceback
        traceback.print_exc()
